"""Word exporter — IC/credit-memo DOCX from a ModelForge workbook.

Generates a structured committee memo with sections:
    1. Executive Summary (auto-generated from key outputs)
    2. Transaction Overview (deal name, version, date, author)
    3. Key Assumptions (top-10 cited drivers)
    4. Sources (the Sources sheet)
    5. Key Outputs (IRR / MOIC / DSCR / Leverage)
    6. QC Pass (12-check gate result)
    7. Recommendation (template — needs author override)

Designed to be a starting point for an IC packet, not a finished doc.
The author edits sections 1 and 7; sections 2-6 are auto-populated.
"""

from __future__ import annotations

from pathlib import Path
from typing import Any

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
except ImportError as e:
    raise ImportError(
        "docx exporter requires `pip install modelforge[export]` "
        "(adds python-docx)."
    ) from e

from openpyxl import load_workbook


NAVY = RGBColor(0x10, 0x2C, 0x57)
BLUE = RGBColor(0x00, 0x00, 0xFF)
GREY = RGBColor(0x55, 0x55, 0x55)
GREEN = RGBColor(0x00, 0x80, 0x00)
RED = RGBColor(0xC0, 0x00, 0x00)


def _read_workbook_meta(xlsx_path: Path) -> dict[str, Any]:
    """Same shape as exporters.pptx._read_workbook_meta — kept independent."""
    wb = load_workbook(xlsx_path, data_only=True)
    meta: dict[str, Any] = {
        "deal_name": xlsx_path.stem,
        "version": "v0.8.9",
        "author": "ModelForge",
        "date": "",
        "assumptions": [],
        "sources": [],
        "qc_checks": [],
        "outputs": {},
    }

    if "Cover" in wb.sheetnames:
        cov = wb["Cover"]
        for row in cov.iter_rows(min_row=1, max_row=15, values_only=True):
            if not row or row[0] is None:
                continue
            label = str(row[0]).strip().lower()
            value = row[1] if len(row) > 1 else None
            if "deal" in label and value:
                meta["deal_name"] = str(value)
            elif "version" in label and value:
                meta["version"] = str(value)
            elif "author" in label and value:
                meta["author"] = str(value)
            elif "date" in label and value:
                meta["date"] = str(value)

    if "Assumptions" in wb.sheetnames:
        asm = wb["Assumptions"]
        for row in asm.iter_rows(min_row=2, max_row=30, values_only=True):
            if not row or row[0] is None:
                continue
            name = str(row[0])
            value = row[1] if len(row) > 1 else None
            source = row[2] if len(row) > 2 else None
            if name and value is not None:
                meta["assumptions"].append({
                    "name": name, "value": value, "source": source,
                })
            if len(meta["assumptions"]) >= 10:
                break

    if "Sources" in wb.sheetnames:
        srcs = wb["Sources"]
        for row in srcs.iter_rows(min_row=2, max_row=50, values_only=True):
            if not row or row[0] is None:
                continue
            sid, doc, page, publisher, *rest = list(row) + [None] * 8
            if sid and doc:
                meta["sources"].append({
                    "id": str(sid),
                    "doc": str(doc),
                    "page": page,
                    "publisher": publisher,
                })

    if "QC" in wb.sheetnames:
        qc = wb["QC"]
        for row in qc.iter_rows(min_row=2, max_row=30, values_only=True):
            if not row or row[0] is None:
                continue
            check, result = row[0], row[1] if len(row) > 1 else None
            meta["qc_checks"].append({"check": str(check), "result": result})

    for sheet_name in ("Returns", "Outputs", "Summary", "KPIs"):
        if sheet_name in wb.sheetnames:
            sh = wb[sheet_name]
            for row in sh.iter_rows(min_row=1, max_row=20, values_only=True):
                if not row or row[0] is None:
                    continue
                label = str(row[0]).strip()
                value = row[1] if len(row) > 1 else None
                if label and value is not None and isinstance(value, (int, float)):
                    meta["outputs"][label] = value
            if meta["outputs"]:
                break

    return meta


def _heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = NAVY


def _kv_table(doc: Document, rows: list[tuple[str, str]], header: tuple[str, str] = ("Item", "Value"),
              value_color: RGBColor = BLUE) -> None:
    if not rows:
        p = doc.add_paragraph("(no entries)")
        for run in p.runs:
            run.font.color.rgb = GREY
        return

    table = doc.add_table(rows=len(rows) + 1, cols=2)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header
    hdr = table.rows[0].cells
    hdr[0].text = header[0]
    hdr[1].text = header[1]
    for c in hdr:
        for p in c.paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = NAVY

    # Rows
    for i, (k, v) in enumerate(rows, start=1):
        row = table.rows[i].cells
        row[0].text = str(k)
        row[1].text = str(v)
        for p in row[1].paragraphs:
            for run in p.runs:
                run.font.color.rgb = value_color


def build_ic_memo(xlsx_path: Path, output_path: Path) -> Path:
    """Build an IC/credit-memo DOCX from a ModelForge workbook.

    Args:
        xlsx_path: Source ModelForge workbook (.xlsx).
        output_path: Output .docx path.

    Returns:
        The output .docx path.
    """
    xlsx_path = Path(xlsx_path)
    output_path = Path(output_path)
    meta = _read_workbook_meta(xlsx_path)

    doc = Document()

    # Title
    title = doc.add_heading(f"Investment Committee Memo — {meta['deal_name']}", level=0)
    for run in title.runs:
        run.font.color.rgb = NAVY

    sub = doc.add_paragraph(
        f"{meta['version']}  ·  {meta['date']}  ·  {meta['author']}  ·  ModelForge"
    )
    for run in sub.runs:
        run.font.color.rgb = GREY
        run.font.size = Pt(10)

    doc.add_paragraph()  # spacing

    # 1. Executive Summary
    _heading(doc, "1. Executive Summary", level=1)
    doc.add_paragraph(
        "[Author override — replace this paragraph with the deal thesis: "
        "what we are recommending, why now, key terms, and the single most "
        "important conviction point. 3-5 sentences, IC-grade prose.]"
    )

    # 2. Transaction Overview
    _heading(doc, "2. Transaction Overview", level=1)
    overview_rows = [
        ("Deal Name", meta["deal_name"]),
        ("Model Version", meta["version"]),
        ("Model Date", meta["date"] or "—"),
        ("Author", meta["author"]),
    ]
    _kv_table(doc, overview_rows, header=("Field", "Value"))

    # 3. Key Assumptions
    _heading(doc, "3. Key Assumptions", level=1)
    doc.add_paragraph(
        "The model is driven by the assumptions below. Each value is sourced "
        "(see Section 4 for full citations). Blue values are hardcoded inputs; "
        "all derived calculations are live-formulated in the underlying workbook."
    )
    a_rows = [
        (a["name"],
         (f"{a['value']:,.2f}" if isinstance(a["value"], (int, float)) else str(a["value"]))
         + (f"  [{a['source']}]" if a.get("source") else ""))
        for a in meta["assumptions"][:10]
    ]
    _kv_table(doc, a_rows, header=("Driver", "Value · Source"), value_color=BLUE)

    # 4. Sources
    _heading(doc, "4. Sources", level=1)
    doc.add_paragraph(
        "Every hardcoded number in the workbook carries a source ID linking "
        "to the documents below. The ModelForge linkage graph (.graph.db) "
        "enables full cell-to-source traceability."
    )
    s_rows = [
        (s["id"], f'{s["doc"]} (p. {s["page"]})' if s.get("page") else s["doc"])
        for s in meta["sources"][:15]
    ]
    _kv_table(doc, s_rows, header=("Source ID", "Document · Page"), value_color=GREEN)

    # 5. Key Outputs
    _heading(doc, "5. Key Outputs", level=1)
    out_rows = [
        (k, f"{v:,.2f}" if isinstance(v, (int, float)) else str(v))
        for k, v in list(meta["outputs"].items())[:10]
    ]
    _kv_table(doc, out_rows, header=("Metric", "BASE Scenario"))

    # 6. QC Gate
    _heading(doc, "6. QC Gate", level=1)
    pass_count = sum(1 for c in meta["qc_checks"] if str(c.get("result", "")).upper() == "PASS")
    total = len(meta["qc_checks"])
    summary = doc.add_paragraph()
    run = summary.add_run(f"{pass_count}/{total} automated checks PASS")
    run.font.bold = True
    run.font.color.rgb = NAVY if pass_count == total else RED

    if meta["qc_checks"]:
        qc_rows = [(c.get("check", "—"), str(c.get("result", "—"))) for c in meta["qc_checks"][:12]]
        _kv_table(doc, qc_rows, header=("Check", "Result"))

    # 7. Recommendation
    _heading(doc, "7. Recommendation", level=1)
    doc.add_paragraph(
        "[Author override — replace with the committee recommendation: "
        "Approve / Approve with Conditions / Decline. Include conditions "
        "and any open items if applicable.]"
    )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path
